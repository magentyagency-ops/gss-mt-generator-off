import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def set_run_font(run, name='Arial', size_pt=11, bold=False, italic=False, color_rgb=(0,0,0)):
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor(*color_rgb)

def add_styled_heading(doc, text, level, space_before=12, space_after=6):
    p = doc.add_heading(level=level)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, name='Arial', size_pt=16, bold=True, color_rgb=(0, 51, 102))
    elif level == 2:
        set_run_font(run, name='Arial', size_pt=14, bold=True, color_rgb=(0, 102, 153))
    else:
        set_run_font(run, name='Arial', size_pt=12, bold=True, color_rgb=(51, 51, 51))
    return p

def add_styled_paragraph(doc, text, space_after=6, bold=False, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, name='Arial', size_pt=11, bold=bold, italic=italic, color_rgb=(51, 51, 51))
    return p

def add_bullet_point(doc, text, space_after=4):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, name='Arial', size_pt=11, color_rgb=(51, 51, 51))
    return p

def create_cctp(dest_path):
    doc = Document()
    
    # Title Page
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(100)
    p_title.paragraph_format.space_after = Pt(20)
    run_title = p_title.add_run("NEOMA BUSINESS SCHOOL\n\nCAHIER DES CLAUSES TECHNIQUES PARTICULIÈRES\n(CCTP)")
    set_run_font(run_title, name='Arial', size_pt=24, bold=True, color_rgb=(0, 51, 102))
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(200)
    run_sub = p_sub.add_run("Prestations de surveillance humaine, de gardiennage et de sécurité incendie\nCampus de Rouen, Reims et Paris")
    set_run_font(run_sub, name='Arial', size_pt=14, italic=True, color_rgb=(102, 102, 102))
    
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_date = p_date.add_run("Version : 1.0 - Juin 2026")
    set_run_font(run_date, name='Arial', size_pt=10, color_rgb=(128, 128, 128))
    
    doc.add_page_break()
    
    # Table of Contents placeholder / Intro
    add_styled_heading(doc, "1. OBJET DU MARCHÉ ET CONTEXTE", level=1)
    add_styled_paragraph(doc, "Le présent marché a pour objet la réalisation de prestations de sécurité privée, de surveillance humaine, de contrôle des accès et de sécurité incendie pour les différents campus de l'école supérieure de commerce NEOMA Business School.")
    add_styled_paragraph(doc, "NEOMA Business School accueille chaque année plus de 10 000 étudiants, enseignants et collaborateurs sur ses trois campus principaux. Le titulaire du marché devra s'assurer de la protection des personnes, des biens immobiliers et mobiliers, ainsi que du maintien de la sécurité incendie 24h/24 et 7j/7 selon les spécificités de chaque site.")
    
    add_styled_heading(doc, "2. DESCRIPTION DES SITES ET BESOINS SPECIFIQUES", level=1)
    
    add_styled_heading(doc, "2.1. Campus de Rouen (Mont-Saint-Aignan)", level=2)
    add_styled_paragraph(doc, "Le campus de Rouen comprend plusieurs bâtiments d'enseignement, des bureaux administratifs, une bibliothèque ainsi que des espaces de restauration et de vie étudiante. Il est situé à Mont-Saint-Aignan.")
    add_styled_paragraph(doc, "Besoins spécifiques :")
    add_bullet_point(doc, "Présence permanente d'un agent de sécurité incendie (SSIAP 1) en période d'ouverture au public.")
    add_bullet_point(doc, "Rondes de surveillance régulières (diurnes et nocturnes) avec contrôle des accès et fermeture des bâtiments.")
    add_bullet_point(doc, "Gestion renforcée lors des événements de l'école (conférences, remises de diplômes, soirées étudiantes).")
    
    add_styled_heading(doc, "2.2. Campus de Reims", level=2)
    add_styled_paragraph(doc, "Le campus de Reims regroupe deux sites géographiques distincts accueillant des salles de cours technologiques, des incubateurs de start-ups et des logements étudiants à proximité.")
    add_styled_paragraph(doc, "Besoins spécifiques :")
    add_bullet_point(doc, "Surveillance humaine et contrôles visuels des sacs / badges à l'entrée principale du campus en période de plan Vigipirate renforcé.")
    add_bullet_point(doc, "Télésurveillance associée à des levées de doute physiques par un agent mobile.")
    
    add_styled_heading(doc, "2.3. Campus de Paris", level=2)
    add_styled_paragraph(doc, "Situé en milieu urbain dense, le campus parisien accueille principalement des étudiants de Master et des cadres en formation continue.")
    add_bullet_point(doc, "Contrôle strict des accès au niveau du hall d'accueil.")
    add_bullet_point(doc, "Fermeture sécurisée des accès en fin de journée et veille technologique des installations d'alarme intrusion.")
    
    add_styled_heading(doc, "3. MODALITÉS D'EXÉCUTION DES PRESTATIONS", level=1)
    add_styled_paragraph(doc, "Les agents déployés devront être titulaires de la carte professionnelle en cours de validité délivrée par le CNAPS et posséder les certifications requises (CQP APS, SST, SSIAP 1 ou 2 selon les postes).")
    add_styled_paragraph(doc, "Le titulaire devra fournir aux agents des équipements de communication modernes (radios VHF, terminaux PTI/DATI) et s'assurer de la traçabilité des rondes de surveillance via un système de main courante électronique accessible en temps réel par NEOMA Business School.")
    
    add_styled_heading(doc, "4. DISPOSITIFS D'ALERTE ET PROCEDURES OPERATIONNELLES", level=1)
    add_styled_paragraph(doc, "En cas de détection d'une anomalie, intrusion, ou départ de feu, les agents devront appliquer immédiatement les consignes de sécurité rédigées conjointement avec la direction des ressources matérielles de NEOMA. Les rapports d'incidents devront être transmis sous 12 heures maximum par voie électronique.")
    
    doc.save(dest_path)

def create_rc(dest_path):
    doc = Document()
    
    # Title Page
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(100)
    p_title.paragraph_format.space_after = Pt(20)
    run_title = p_title.add_run("NEOMA BUSINESS SCHOOL\n\nRÈGLEMENT DE LA CONSULTATION\n(RC)")
    set_run_font(run_title, name='Arial', size_pt=24, bold=True, color_rgb=(0, 51, 102))
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(200)
    run_sub = p_sub.add_run("Prestations de surveillance humaine et de sécurité incendie\nProcédure Adaptée Ouverte")
    set_run_font(run_sub, name='Arial', size_pt=14, italic=True, color_rgb=(102, 102, 102))
    
    doc.add_page_break()
    
    add_styled_heading(doc, "1. CONDITIONS DE LA CONSULTATION", level=1)
    add_styled_paragraph(doc, "La présente consultation est lancée selon une procédure adaptée ouverte, en application des règles internes de NEOMA Business School. L'école se réserve la possibilité de négocier avec les trois candidats ayant présenté les meilleures offres à l'issue d'une première analyse.")
    
    add_styled_heading(doc, "2. CALENDRIER DE LA CONSULTATION", level=1)
    add_bullet_point(doc, "Date de publication de l'avis : 22 Juin 2026")
    add_bullet_point(doc, "Date limite de réception des offres : 15 Septembre 2026 à 12h00")
    add_bullet_point(doc, "Date prévisionnelle de signature du marché : 15 Octobre 2026")
    add_bullet_point(doc, "Date de début d'exécution : 1er Novembre 2026")
    
    add_styled_heading(doc, "3. DOSSIER DE CONSULTATION DES ENTREPRISES (DCE)", level=1)
    add_styled_paragraph(doc, "Le dossier mis à disposition des candidats comprend :")
    add_bullet_point(doc, "Le présent Règlement de la Consultation (RC)")
    add_bullet_point(doc, "Le Cahier des Clauses Techniques Particulières (CCTP) et ses annexes")
    add_bullet_point(doc, "Le Cahier des Clauses Administratives Particulières (CCAP)")
    add_bullet_point(doc, "L'Acte d'Engagement (AE)")
    
    add_styled_heading(doc, "4. CRITÈRES D'ATTRIBUTION DES OFFRES", level=1)
    add_styled_paragraph(doc, "Le jugement des offres sera basé sur les critères pondérés suivants :")
    
    add_styled_heading(doc, "4.1. Valeur Technique (Mémoire Technique) - 60%", level=2)
    add_styled_paragraph(doc, "La valeur technique sera évaluée sur la base du mémoire technique remis par le candidat, selon la grille suivante :")
    add_bullet_point(doc, "Qualité de l'organisation humaine et méthodologie de planification (20 points)")
    add_bullet_point(doc, "Moyens matériels et technologiques mis en œuvre (15 points)")
    add_bullet_point(doc, "Procédures d'urgence, de contrôle qualité et de réactivité (15 points)")
    add_bullet_point(doc, "Politique RSE et engagements environnementaux de la structure (10 points)")
    
    add_styled_heading(doc, "4.2. Prix des prestations - 40%", level=2)
    add_styled_paragraph(doc, "L'offre financière sera évaluée à partir du montant total annuel théorique simulé à l'aide des bordereaux de prix unitaires fournis par le candidat.")
    
    doc.save(dest_path)

def create_ccap(dest_path):
    doc = Document()
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(100)
    p_title.paragraph_format.space_after = Pt(20)
    run_title = p_title.add_run("NEOMA BUSINESS SCHOOL\n\nCAHIER DES CLAUSES ADMINISTRATIVES PARTICULIÈRES\n(CCAP)")
    set_run_font(run_title, name='Arial', size_pt=20, bold=True, color_rgb=(0, 51, 102))
    doc.add_page_break()
    
    add_styled_heading(doc, "ARTICLE 1 - PIÈCES CONSTITUTIVES DU MARCHÉ", level=1)
    add_styled_paragraph(doc, "Les pièces constitutives du marché sont, par ordre de priorité décroissant :")
    add_bullet_point(doc, "L'Acte d'Engagement dûment complété et signé par les parties.")
    add_bullet_point(doc, "Le Cahier des Clauses Techniques Particulières (CCTP) et ses annexes.")
    add_bullet_point(doc, "Le présent Cahier des Clauses Administratives Particulières (CCAP).")
    add_bullet_point(doc, "Le mémoire technique et l'offre financière du titulaire retenu.")
    
    add_styled_heading(doc, "ARTICLE 2 - DURÉE ET RECONDUCTION DU MARCHÉ", level=1)
    add_styled_paragraph(doc, "Le marché est conclu pour une période initiale de un (1) an à compter du 1er Novembre 2026. Il pourra être reconduit expressément trois (3) fois pour des périodes successives de un an, sans que sa durée totale ne puisse excéder quatre (4) ans.")
    
    add_styled_heading(doc, "ARTICLE 3 - PENALITES ET CONTROLES", level=1)
    add_styled_paragraph(doc, "Des pénalités forfaitaires de 150 € par heure de retard constatée ou d'absence d'un agent planifié seront appliquées de plein droit si le titulaire ne remédie pas à la carence dans un délai d'une heure. Les contrôles pourront être réalisés de manière inopinée par les équipes de la Direction Générale de NEOMA.")
    
    doc.save(dest_path)

def create_ae(dest_path):
    doc = Document()
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(100)
    p_title.paragraph_format.space_after = Pt(20)
    run_title = p_title.add_run("NEOMA BUSINESS SCHOOL\n\nACTE D'ENGAGEMENT (AE)")
    set_run_font(run_title, name='Arial', size_pt=20, bold=True, color_rgb=(0, 51, 102))
    doc.add_page_break()
    
    add_styled_heading(doc, "1. IDENTIFICATION DU CANDIDAT", level=1)
    add_styled_paragraph(doc, "Dénomination sociale : __________________________________")
    add_styled_paragraph(doc, "Adresse du siège social : ________________________________")
    add_styled_paragraph(doc, "Représentée par (Nom, Prénom et qualité) : _________________")
    
    add_styled_heading(doc, "2. ENGAGEMENT DU CANDIDAT", level=1)
    add_styled_paragraph(doc, "Après avoir pris connaissance du dossier de consultation (RC, CCTP, CCAP) pour les prestations de sécurité des campus NEOMA Business School, le soussigné s'engage à exécuter lesdites prestations conformément aux clauses et conditions définies, moyennant les tarifs portés au Bordereau des Prix Unitaires.")
    
    add_styled_heading(doc, "3. SIGNATURE DE L'OFFRE", level=1)
    add_styled_paragraph(doc, "Fait à : __________________ Le : _________________")
    add_styled_paragraph(doc, "Signature et cachet commercial du candidat :")
    
    doc.save(dest_path)

if __name__ == '__main__':
    desktop_dir = '/Users/clarencegomis/Desktop'
    target_dir = os.path.join(desktop_dir, 'DCE_NEOMA')
    if not os.path.exists(target_dir):
        os.makedirs(target_dir)
        
    create_cctp(os.path.join(target_dir, '1-CCTP_Securite_NEOMA.docx'))
    create_rc(os.path.join(target_dir, '2-RC_Securite_NEOMA.docx'))
    create_ccap(os.path.join(target_dir, '3-CCAP_Securite_NEOMA.docx'))
    create_ae(os.path.join(target_dir, '4-Acte_Engagement_NEOMA.docx'))
    print(f"DCE NEOMA créé avec succès dans {target_dir}")

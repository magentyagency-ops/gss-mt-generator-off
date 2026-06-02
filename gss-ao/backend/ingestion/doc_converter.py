"""Conversion et extraction de texte multi-formats — fondation des parseurs.

Responsabilités :
  - détecter le binaire LibreOffice (`soffice`) — PATH ou emplacement macOS ;
  - convertir un `.doc` legacy en `.docx` via LibreOffice headless ;
  - charger un `.docx` (python-docx) — natif ou issu d'une conversion ;
  - extraire le texte d'un PDF (PyMuPDF) — fallback OCR OPTIONNEL hors périmètre.

Principe de sûreté : si LibreOffice est requis mais absent, on lève
`LibreOfficeNotFoundError` avec un message actionnable — JAMAIS un résultat
silencieusement faux (exigence brief §12 : revue humaine, pas d'illusion).
"""

from __future__ import annotations

import shutil
import subprocess
import tempfile
from pathlib import Path

from backend.core.config import Settings, get_settings

# Emplacements macOS usuels de LibreOffice (hors PATH).
_MACOS_SOFFICE_CANDIDATES = (
    "/Applications/LibreOffice.app/Contents/MacOS/soffice",
    "/opt/homebrew/bin/soffice",
    "/usr/local/bin/soffice",
)


class DocConverterError(RuntimeError):
    """Erreur générique de conversion/extraction."""


class LibreOfficeNotFoundError(DocConverterError):
    """LibreOffice est requis (conversion .doc) mais introuvable."""


def find_soffice(settings: Settings | None = None) -> str | None:
    """Localise le binaire `soffice`.

    Ordre : `SOFFICE_BIN` (config) > PATH (`soffice`/`libreoffice`) > candidats
    macOS. Retourne le chemin absolu ou None si introuvable.
    """
    settings = settings or get_settings()

    if settings.soffice_bin:
        p = Path(settings.soffice_bin)
        if p.exists():
            return str(p)

    for name in ("soffice", "libreoffice"):
        found = shutil.which(name)
        if found:
            return found

    for candidate in _MACOS_SOFFICE_CANDIDATES:
        if Path(candidate).exists():
            return candidate

    return None


def soffice_available(settings: Settings | None = None) -> bool:
    """True si LibreOffice est utilisable."""
    return find_soffice(settings) is not None


def find_textutil() -> str | None:
    """Localise `textutil` (outil natif macOS de conversion de documents).

    Permet d'extraire le texte d'un `.doc` legacy SANS LibreOffice (macOS only).
    """
    return shutil.which("textutil") or (
        "/usr/bin/textutil" if Path("/usr/bin/textutil").exists() else None
    )


def convert_doc_to_text_textutil(src: Path, *, timeout: int = 120) -> str:
    """Extrait le texte d'un `.doc`/`.rtf`/`.docx` via `textutil` (macOS).

    Raises:
        DocConverterError: textutil absent (non-macOS) ou échec de conversion.
    """
    src = Path(src)
    if not src.exists():
        raise DocConverterError(f"Fichier introuvable : {src}")
    textutil = find_textutil()
    if textutil is None:
        raise DocConverterError("textutil introuvable (disponible uniquement sur macOS).")
    try:
        proc = subprocess.run(
            [textutil, "-convert", "txt", "-stdout", str(src)],
            capture_output=True,
            text=True,
            timeout=timeout,
            check=False,
        )
    except subprocess.TimeoutExpired as exc:
        raise DocConverterError(
            f"Conversion textutil expirée après {timeout}s : {src.name}"
        ) from exc
    if proc.returncode != 0:
        raise DocConverterError(
            f"Échec textutil ({src.name}) : {proc.stderr.strip()[:300]}"
        )
    return proc.stdout


def convert_doc_to_docx(
    src: Path,
    out_dir: Path | None = None,
    settings: Settings | None = None,
    *,
    timeout: int = 120,
) -> Path:
    """Convertit un `.doc` (ou autre format Word) en `.docx` via LibreOffice.

    Args:
        src: fichier source (.doc / .rtf / .odt ...).
        out_dir: dossier de sortie ; défaut = dossier temporaire.
        timeout: secondes max pour le sous-processus soffice.

    Returns:
        Chemin du `.docx` produit.

    Raises:
        LibreOfficeNotFoundError: binaire soffice introuvable.
        DocConverterError: échec de conversion / fichier de sortie absent.
    """
    src = Path(src)
    if not src.exists():
        raise DocConverterError(f"Fichier introuvable : {src}")

    soffice = find_soffice(settings)
    if soffice is None:
        raise LibreOfficeNotFoundError(
            "LibreOffice (soffice) est requis pour convertir les .doc legacy mais "
            "n'a pas été trouvé. Installez-le puis réessayez, ou renseignez "
            "SOFFICE_BIN dans .env. Voir DECISIONS.md §D1."
        )

    out_dir = Path(out_dir) if out_dir else Path(tempfile.mkdtemp(prefix="gss_doc_"))
    out_dir.mkdir(parents=True, exist_ok=True)

    # LibreOffice headless : un seul profil utilisateur jetable pour éviter les
    # conflits si une instance GUI tourne déjà.
    with tempfile.TemporaryDirectory(prefix="gss_lo_profile_") as profile:
        cmd = [
            soffice,
            "--headless",
            "--norestore",
            f"-env:UserInstallation=file://{profile}",
            "--convert-to",
            "docx",
            "--outdir",
            str(out_dir),
            str(src),
        ]
        try:
            proc = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=timeout,
                check=False,
            )
        except subprocess.TimeoutExpired as exc:
            raise DocConverterError(
                f"Conversion LibreOffice expirée après {timeout}s : {src.name}"
            ) from exc

    produced = out_dir / (src.stem + ".docx")
    if proc.returncode != 0 or not produced.exists():
        raise DocConverterError(
            f"Échec conversion LibreOffice ({src.name}). "
            f"code={proc.returncode} stderr={proc.stderr.strip()[:500]}"
        )
    return produced


def load_docx_text(path: Path) -> str:
    """Extrait le texte brut d'un `.docx` (paragraphes + tables), séparé par \\n."""
    from docx import Document

    doc = Document(str(path))
    lines: list[str] = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            lines.append(" | ".join(cells))
    return "\n".join(lines)


def extract_pdf_text(path: Path) -> str:
    """Extrait le texte d'un PDF via PyMuPDF (pages séparées par form feed).

    Note : OCR (PDF scannés) est un fallback OPTIONNEL hors périmètre itération 1
    (OCR_ENABLED=false par défaut).
    """
    import fitz  # PyMuPDF

    parts: list[str] = []
    with fitz.open(str(path)) as doc:
        for page in doc:
            parts.append(page.get_text("text"))
    return "\f".join(parts)


def extract_doc_text(src: Path, settings: Settings | None = None):
    """Extrait le texte d'un `.doc` legacy, retourne (texte, méthode_extraction).

    Ordre de préférence :
      1. LibreOffice (conversion .docx -> python-docx) : préserve les tableaux ;
      2. textutil (natif macOS) : zéro dépendance, texte propre.

    Raises:
        LibreOfficeNotFoundError: aucune voie de conversion disponible.
    """
    from backend.schemas.common import ExtractionMethod

    src = Path(src)
    if soffice_available(settings):
        converted = convert_doc_to_docx(src, settings=settings)
        return load_docx_text(converted), ExtractionMethod.DOC_LIBREOFFICE
    if find_textutil() is not None:
        return convert_doc_to_text_textutil(src), ExtractionMethod.DOC_TEXTUTIL
    raise LibreOfficeNotFoundError(
        "Aucune voie de conversion .doc disponible : ni LibreOffice (soffice) ni "
        "textutil (macOS). Installez LibreOffice ou renseignez SOFFICE_BIN. "
        "Voir DECISIONS.md §D1."
    )


def extract_text(path: Path, settings: Settings | None = None) -> str:
    """Extrait le texte de n'importe quel format supporté (dispatch par extension).

    - `.docx` -> python-docx
    - `.doc`  -> LibreOffice si dispo, sinon textutil (macOS)
    - `.pdf`  -> PyMuPDF
    """
    path = Path(path)
    ext = path.suffix.lower()
    if ext == ".docx":
        return load_docx_text(path)
    if ext == ".doc":
        text, _ = extract_doc_text(path, settings=settings)
        return text
    if ext == ".pdf":
        return extract_pdf_text(path)
    raise DocConverterError(f"Format non supporté : {ext} ({path.name})")

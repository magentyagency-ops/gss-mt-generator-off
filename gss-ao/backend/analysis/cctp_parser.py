"""Module B — Parseur CCTP (Cahier des Clauses Techniques Particulières).

OPÉRATIONNEL (brief §14.3). Cible : `4-CCTP 2026-08.docx` (Université de Rouen).

Approche :
  1. Reconstruire l'arborescence hiérarchique à partir des styles de titre
     (Heading 1/2/3 — python-docx traduit le styleId 'Titre1' en 'Heading 1').
  2. Rattacher le texte courant à la section la plus profonde.
  3. Dériver, par motifs sur les titres :
       - les prestations (base / supplémentaire / télésécurité) + lot + campus ;
       - les exigences agents (qualification / tenue / équipement / comportement
         / véhicule) ;
       - la reprise du personnel et les contraintes de site (ex. ZRR).

Robuste à la variabilité : on ne hardcode pas des numéros de section, on
reconnaît des motifs lexicaux (le format CCTP varie d'un AO à l'autre).
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph

from backend.schemas.cctp import (
    CategorieExigence,
    CCTPDocument,
    ExigenceAgent,
    Prestation,
    SectionCCTP,
    TypePrestation,
)
from backend.schemas.common import ExtractionMethod, SourceMeta

# --- détection du niveau de titre -------------------------------------------
# python-docx expose 'Heading N' ; on accepte aussi 'Titre N' (FR) par sécurité.
_HEADING_RE = re.compile(r"^(?:Heading|Titre)\s*(\d+)$", re.IGNORECASE)


def _heading_level(p: Paragraph) -> int | None:
    """Retourne le niveau de titre (1..n) ou None si le paragraphe n'en est pas un."""
    name = (p.style.name or "") if p.style else ""
    m = _HEADING_RE.match(name.strip())
    return int(m.group(1)) if m else None


# --- motifs lexicaux ---------------------------------------------------------
_LOT_RE = re.compile(r"lot\s*(\d+)", re.IGNORECASE)
_DEPT_TO_LOT = {"76": 1, "27": 2}  # Seine-Maritime -> lot 1, Eure -> lot 2

_KW_TELESECU = ("télé-sécurité", "télésécurité", "télé sécurité", "telesecurite")
_KW_SUPPL = ("supplémentaire", "supplementaire", "à la demande", "a la demande")
_KW_BASE = ("de base", "« de base »", "prestations de base")

# Qualifications réglementaires à repérer dans le texte (brief §5.3).
_QUALIF_TOKENS = (
    "SSIAP2", "SSIAP 2", "SSIAP1", "SSIAP 1", "APS", "ADS", "AMC",
    "CNAPS", "carte professionnelle", "recyclage", "H0B0", "HOB0", "SST", "PSC1",
)

# Sous-titres -> catégorie d'exigence agent.
_EXIGENCE_TITLE_MAP = (
    ("qualification", CategorieExigence.QUALIFICATION),
    ("recyclage", CategorieExigence.QUALIFICATION),
    ("tenue", CategorieExigence.TENUE),
    ("équipement", CategorieExigence.EQUIPEMENT),
    ("equipement", CategorieExigence.EQUIPEMENT),
    ("véhicule", CategorieExigence.VEHICULE),
    ("vehicule", CategorieExigence.VEHICULE),
    ("comportement", CategorieExigence.COMPORTEMENT),
    ("présentation", CategorieExigence.COMPORTEMENT),
    ("presentation", CategorieExigence.COMPORTEMENT),
)


def _classify_prestation_type(title: str) -> TypePrestation | None:
    low = title.lower()
    if any(k in low for k in _KW_TELESECU):
        return TypePrestation.TELESECURITE
    if any(k in low for k in _KW_SUPPL):
        return TypePrestation.SUPPLEMENTAIRE
    if any(k in low for k in _KW_BASE):
        return TypePrestation.BASE
    return None


def _detect_lot(title: str) -> int | None:
    m = _LOT_RE.search(title)
    if m:
        return int(m.group(1))
    for dept, lot in _DEPT_TO_LOT.items():
        if dept in title:
            return lot
    return None


def _build_tree(doc: Document) -> tuple[list[SectionCCTP], list[SectionCCTP]]:
    """Construit l'arborescence et retourne (racines, liste_à_plat_dans_l_ordre)."""
    roots: list[SectionCCTP] = []
    flat: list[SectionCCTP] = []
    # pile de (niveau, section)
    stack: list[tuple[int, SectionCCTP]] = []

    for p in doc.paragraphs:
        lvl = _heading_level(p)
        text = p.text.strip()
        if lvl is not None and text:
            section = SectionCCTP(niveau=lvl, titre=text)
            flat.append(section)
            # dépiler jusqu'à un parent de niveau < lvl
            while stack and stack[-1][0] >= lvl:
                stack.pop()
            if stack:
                stack[-1][1].enfants.append(section)
            else:
                roots.append(section)
            stack.append((lvl, section))
        elif text and stack:
            # texte courant rattaché à la section ouverte la plus profonde
            cur = stack[-1][1]
            cur.texte = (cur.texte + "\n" + text).strip() if cur.texte else text

    return roots, flat


def _full_text_of(section: SectionCCTP) -> str:
    """Concatène le texte d'une section et de ses descendants."""
    parts = [section.texte] if section.texte else []
    for child in section.enfants:
        parts.append(_full_text_of(child))
    return "\n".join(parts).strip()


# Titres de niveau 1 qui NE sont PAS une branche de prestations
# (on n'y cherche pas de prestations, pour éviter les faux positifs).
_NON_PRESTATION_L1 = (
    "obligations du titulaire",
    "dispositions particulieres",
    "dispositions particulières",
    "accueil de stagiaires",
    "contrôle interne",
    "controle interne",
)


def _is_prestation_branch(l1_title: str | None) -> bool:
    if l1_title is None:
        return True
    low = l1_title.lower()
    return not any(k in low for k in _NON_PRESTATION_L1)


def _extract_prestations(flat: list[SectionCCTP]) -> list[Prestation]:
    """Dérive les prestations à partir des titres de section.

    Le lot courant est suivi via les titres de niveau 1 (ARTICLE ... lot N).
    Les branches non liées aux prestations (obligations, dispositions agents,
    accueil stagiaires, contrôle) sont ignorées pour limiter les faux positifs.
    """
    prestations: list[Prestation] = []
    current_lot: int | None = None
    current_campus: str | None = None
    current_l1: str | None = None

    for sec in flat:
        if sec.niveau == 1:
            current_l1 = sec.titre
            current_lot = _detect_lot(sec.titre)
            current_campus = None
        # un L3 sous un "détail par campus" donne souvent le campus
        if sec.niveau == 3 and "campus" in sec.titre.lower():
            current_campus = sec.titre

        if not _is_prestation_branch(current_l1):
            continue

        ptype = _classify_prestation_type(sec.titre)
        if ptype is not None:
            lot = current_lot
            if ptype is TypePrestation.TELESECURITE and lot is None:
                lot = 3
            prestations.append(
                Prestation(
                    type=ptype,
                    lot=lot,
                    campus=current_campus if sec.niveau >= 3 else None,
                    description=sec.titre,
                    ref_section=sec.titre,
                )
            )
    return prestations


def _extract_exigences_agents(flat: list[SectionCCTP]) -> list[ExigenceAgent]:
    """Extrait les exigences agents : (a) sous-sections dédiées repérées par
    titre, (b) qualifications réglementaires détectées dans le texte global."""
    exigences: list[ExigenceAgent] = []
    seen: set[tuple[str, str]] = set()

    # (a) sections dont le titre dénote une exigence
    for sec in flat:
        low = sec.titre.lower()
        for token, categorie in _EXIGENCE_TITLE_MAP:
            if token in low:
                key = (categorie.value, sec.titre)
                if key not in seen:
                    seen.add(key)
                    exigences.append(
                        ExigenceAgent(
                            categorie=categorie,
                            libelle=sec.titre,
                            valeur=(sec.texte[:500] or None),
                            ref_section=sec.titre,
                        )
                    )
                break

    # (b) qualifications réglementaires citées dans le corps du document
    full = "\n".join(_full_text_of(s) for s in flat if s.niveau == 1)
    full_plus_titles = full + "\n" + "\n".join(s.titre for s in flat)
    for token in _QUALIF_TOKENS:
        if re.search(re.escape(token), full_plus_titles, re.IGNORECASE):
            norm = token.upper().replace(" ", "")
            key = (CategorieExigence.QUALIFICATION.value, norm)
            if key not in seen:
                seen.add(key)
                exigences.append(
                    ExigenceAgent(
                        categorie=CategorieExigence.QUALIFICATION,
                        libelle=f"Qualification requise : {token}",
                        valeur=token,
                        ref_section=None,
                    )
                )
    return exigences


def _detect_reprise_personnel(flat: list[SectionCCTP]) -> bool:
    return any("reprise du personnel" in s.titre.lower() for s in flat)


def _detect_contraintes_site(flat: list[SectionCCTP]) -> list[str]:
    contraintes: list[str] = []
    for sec in flat:
        low = sec.titre.lower()
        if "zrr" in low or "zone" in low and "restrict" in low:
            contraintes.append(sec.titre)
        if "filtrage" in low:
            contraintes.append(sec.titre)
    return contraintes


def _detect_objet(doc: Document, flat: list[SectionCCTP]) -> str | None:
    # 1) style "Objet" explicite (s'il est renseigné)
    for p in doc.paragraphs:
        if p.style and p.style.name == "Objet" and p.text.strip():
            return p.text.strip()
    # 2) paragraphe suivant "Ayant pour objet :"
    paras = [p.text.strip() for p in doc.paragraphs]
    for i, txt in enumerate(paras):
        if "ayant pour objet" in txt.lower():
            for nxt in paras[i + 1 :]:
                if nxt:
                    return nxt
    # 3) repli : premier titre de l'arborescence
    return flat[0].titre if flat else None


def parse_cctp(path: Path, settings=None) -> CCTPDocument:
    """Parse un CCTP `.docx` et retourne un `CCTPDocument` structuré.

    Args:
        path: chemin du `.docx` (un `.doc` doit être converti en amont via
            backend.ingestion.doc_converter.convert_doc_to_docx).
    """
    path = Path(path)
    warnings: list[str] = []
    if path.suffix.lower() != ".docx":
        raise ValueError(
            f"parse_cctp attend un .docx (reçu {path.suffix}). "
            "Convertir d'abord via doc_converter.convert_doc_to_docx."
        )

    doc = Document(str(path))
    roots, flat = _build_tree(doc)
    if not flat:
        warnings.append("Aucun titre (Heading/Titre) détecté — structure inattendue.")

    prestations = _extract_prestations(flat)
    exigences = _extract_exigences_agents(flat)
    if not prestations:
        warnings.append("Aucune prestation dérivée des titres.")
    if not exigences:
        warnings.append("Aucune exigence agent détectée.")

    return CCTPDocument(
        objet=_detect_objet(doc, flat),
        arborescence=roots,
        prestations=prestations,
        exigences_agents=exigences,
        contraintes_site=_detect_contraintes_site(flat),
        reprise_personnel=_detect_reprise_personnel(flat),
        source=SourceMeta(
            fichier=path.name,
            methode_extraction=ExtractionMethod.DOCX_NATIVE,
            warnings=warnings,
        ),
    )
